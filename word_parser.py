#!/usr/bin/env python3
import os
from docx import Document
import argparse
from datetime import datetime

def extract_text_from_word(docx_path):
    """
    从Word文档中提取文本内容
    :param docx_path: Word文档路径
    :return: 提取的文本内容
    """
    try:
        # 打开Word文档
        doc = Document(docx_path)
        
        # 提取所有段落的文本
        full_text = []
        
        # 提取标题
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    
    except Exception as e:
        print(f"处理文档时出错: {e}")
        return None

def save_text_to_file(text, output_path):
    """
    将文本保存到文件
    :param text: 要保存的文本内容
    :param output_path: 输出文件路径
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"文本已成功保存到: {output_path}")
    except Exception as e:
        print(f"保存文件时出错: {e}")

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='Word文档解析工具')
    parser.add_argument('input_file', help='输入的Word文档路径')
    parser.add_argument('-o', '--output', help='输出的文本文件路径（可选）')
    
    args = parser.parse_args()
    
    # 检查输入文件是否存在
    if not os.path.exists(args.input_file):
        print(f"错误：找不到输入文件 {args.input_file}")
        return
    
    # 如果没有指定输出文件，则使用默认名称
    if not args.output:
        input_filename = os.path.splitext(os.path.basename(args.input_file))[0]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        args.output = f"{input_filename}_{timestamp}.txt"
    
    # 提取文本
    text_content = extract_text_from_word(args.input_file)
    
    if text_content:
        # 保存文本
        save_text_to_file(text_content, args.output)
    else:
        print("未能成功提取文本内容")

if __name__ == "__main__":
    main() 